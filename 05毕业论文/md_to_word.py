# -*- coding: utf-8 -*-
"""
赣东学院本科毕业论文 md → Word 转换脚本
依赖：pip install python-docx
用法：python md_to_word.py
输出：与本脚本同目录下生成 output.docx
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────
# 路径配置（脚本与 md 文件在同一目录）
# ─────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MD_FILE    = os.path.join(SCRIPT_DIR, "20222042236-吴兆国-基于FreeRTOS的智能居家安防系统.md")
OUT_FILE   = os.path.join(SCRIPT_DIR, "20222042236-吴兆国-基于FreeRTOS的智能居家安防系统.docx")

# ─────────────────────────────────────────────
# 颜色常量
# ─────────────────────────────────────────────
COLOR_BLACK      = RGBColor(0x00, 0x00, 0x00)
COLOR_CODE_BG    = RGBColor(0xF5, 0xF5, 0xF5)   # 代码块背景：浅灰
COLOR_CODE_BORDER= RGBColor(0xCC, 0xCC, 0xCC)   # 代码块边框：灰色

# ─────────────────────────────────────────────
# 辅助：设置段落行距
# ─────────────────────────────────────────────
def set_line_spacing(para, rule=WD_LINE_SPACING.MULTIPLE, val=1.25):
    pf = para.paragraph_format
    pf.line_spacing_rule = rule
    pf.line_spacing = val

# ─────────────────────────────────────────────
# 辅助：设置段落缩进（首行缩进，单位 Pt）
# ─────────────────────────────────────────────
def set_first_indent(para, pt_val):
    para.paragraph_format.first_line_indent = Pt(pt_val)

# ─────────────────────────────────────────────
# 辅助：给段落添加底纹（用于代码块背景）
# ─────────────────────────────────────────────
def set_para_shading(para, fill_hex="F5F5F5"):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)

# ─────────────────────────────────────────────
# 辅助：给段落加边框（上下左右，用于代码块）
# ─────────────────────────────────────────────
def set_para_border(para, color="CCCCCC", sz="4"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        bd = OxmlElement(f'w:{side}')
        bd.set(qn('w:val'), 'single')
        bd.set(qn('w:sz'), sz)
        bd.set(qn('w:space'), '4')
        bd.set(qn('w:color'), color)
        pBdr.append(bd)
    pPr.append(pBdr)

# ─────────────────────────────────────────────
# 辅助：解析行内 **bold** 和普通文本，返回 [(text, bold)]
# ─────────────────────────────────────────────
def parse_inline(text):
    """把 **粗体** 拆成 [(str, is_bold), ...]"""
    parts = []
    pattern = re.compile(r'\*\*(.+?)\*\*')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False))
        parts.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False))
    return parts if parts else [(text, False)]

# ─────────────────────────────────────────────
# 辅助：向段落添加带格式的 run
# ─────────────────────────────────────────────
def add_run(para, text, bold=False, font_name_cn="宋体", font_name_en="Times New Roman",
            font_size_pt=12, color=COLOR_BLACK, italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    run.font.size = Pt(font_size_pt)
    # 中文字体
    run.font.name = font_name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
    return run

# ─────────────────────────────────────────────
# 辅助：向段落添加解析了 **bold** 的文本
# ─────────────────────────────────────────────
def add_inline_text(para, text, font_name_cn="宋体", font_name_en="Times New Roman",
                    font_size_pt=12, base_bold=False, color=COLOR_BLACK):
    for seg, is_bold in parse_inline(text):
        add_run(para, seg, bold=(base_bold or is_bold),
                font_name_cn=font_name_cn, font_name_en=font_name_en,
                font_size_pt=font_size_pt, color=color)

# ─────────────────────────────────────────────
# 辅助：三线表格样式
# ─────────────────────────────────────────────
def set_three_line_table(table):
    """上下粗线(1.5磅)，中间细线(0.5磅)，无竖线"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def make_border(val, sz, color="000000"):
        b = OxmlElement('w:' + val)
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        return b

    rows = table.rows
    n = len(rows)
    for i, row in enumerate(rows):
        for j, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')

            # 顶线
            if i == 0:
                tcBorders.append(make_border('top', 12))       # 1.5磅 = 12 eighths
            else:
                tcBorders.append(make_border('top', 4))        # 0.5磅 = 4 eighths

            # 底线
            if i == n - 1:
                tcBorders.append(make_border('bottom', 12))
            else:
                tcBorders.append(make_border('bottom', 4))

            # 无左右竖线
            for side in ('left', 'right'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'none')
                tcBorders.append(b)

            # 无内部竖线
            for side in ('insideH', 'insideV'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'none')
                tcBorders.append(b)

            tcPr.append(tcBorders)

# ─────────────────────────────────────────────
# 页面设置：A4，页边距
# ─────────────────────────────────────────────
def setup_page(doc):
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.8)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.header_distance = Cm(1.6)
    section.footer_distance = Cm(1.5)

# ─────────────────────────────────────────────
# 添加页眉
# ─────────────────────────────────────────────
def add_header(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run("毕业论文（设计）")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)   # 5号

# ─────────────────────────────────────────────
# 主转换逻辑
# ─────────────────────────────────────────────
def convert(md_path, out_path):
    with open(md_path, encoding='utf-8') as f:
        raw = f.read()

    doc = Document()
    setup_page(doc)
    add_header(doc)

    # 删除默认空段落
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    lines = raw.splitlines()
    i = 0
    total = len(lines)

    def peek(offset=1):
        idx = i + offset
        return lines[idx] if idx < total else ""

    while i < total:
        line = lines[i]

        # ── 空行 / 分隔线 ──────────────────────────────
        if line.strip() == "" or line.strip() == "---":
            i += 1
            continue

        # ── 代码块 ─────────────────────────────────────
        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < total and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结束 ```
            for cl in code_lines:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_line_spacing(p, WD_LINE_SPACING.MULTIPLE, 1.0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                p.paragraph_format.left_indent  = Cm(0.5)
                set_para_shading(p, "F5F5F5")
                set_para_border(p, "CCCCCC", "4")
                run = p.add_run(cl if cl else " ")
                run.font.name = "Courier New"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            continue

        # ── Markdown 表格 ──────────────────────────────
        if line.strip().startswith("|"):
            table_lines = []
            while i < total and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            # 过滤分隔行
            data_rows = [r for r in table_lines
                         if not re.match(r'^\s*\|[\s\-\|:]+\|\s*$', r)]
            if not data_rows:
                continue
            # 解析单元格
            def parse_row(r):
                cells = [c.strip() for c in r.strip().strip('|').split('|')]
                return cells

            rows_data = [parse_row(r) for r in data_rows]
            col_count = max(len(r) for r in rows_data)
            # 补齐列数
            rows_data = [r + [''] * (col_count - len(r)) for r in rows_data]

            tbl = doc.add_table(rows=len(rows_data), cols=col_count)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.style = 'Table Grid'

            for ri, row_cells in enumerate(rows_data):
                for ci, cell_text in enumerate(row_cells):
                    cell = tbl.cell(ri, ci)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cp = cell.paragraphs[0]
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_line_spacing(cp, WD_LINE_SPACING.MULTIPLE, 1.25)
                    is_header = (ri == 0)
                    add_inline_text(cp, cell_text,
                                    font_name_cn="宋体",
                                    font_name_en="Times New Roman",
                                    font_size_pt=10.5,   # 五号
                                    base_bold=is_header)
            set_three_line_table(tbl)
            # 表格后空一行
            doc.add_paragraph()
            continue

        # ── 标题识别 ───────────────────────────────────
        # 章标题：纯数字开头 "1. xxx" 或 "绪论" "结论" "参考文献" "致谢" "附录"
        # 节标题：1.1 xxx
        # 条标题：1.1.1 xxx

        stripped = line.strip()

        # 去掉 Markdown # 前缀（兼容）
        if stripped.startswith('#'):
            stripped = re.sub(r'^#+\s*', '', stripped)

        # 特殊独立章：绪论/结论/参考文献/致谢/附录
        special_chapters = re.match(
            r'^(绪论|结\s*论|参考文献|致\s*谢|附\s*录)(.*)', stripped)

        # 章：1. xxx  或  1 xxx（数字+点/空格开头，后面不再有点）
        chapter_match = re.match(r'^(\d+)[\.。]\s+(.+)', stripped)

        # 节：1.1 xxx
        section_match = re.match(r'^(\d+\.\d+)\s+(.+)', stripped)

        # 条：1.1.1 xxx
        subsec_match = re.match(r'^(\d+\.\d+\.\d+)\s+(.+)', stripped)

        # 附录编号：附录1 / 附录一
        appendix_match = re.match(r'^(附录\s*[\d一二三四五六七八九十]+)(.*)', stripped)

        if special_chapters:
            title_text = special_chapters.group(1) + special_chapters.group(2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            set_line_spacing(p, WD_LINE_SPACING.MULTIPLE, 1.25)
            add_run(p, title_text, bold=True,
                    font_name_cn="黑体", font_name_en="Times New Roman",
                    font_size_pt=15)   # 小三号 = 15pt
            i += 1
            continue

        if appendix_match:
            title_text = appendix_match.group(1) + appendix_match.group(2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.74)  # 空两格
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            set_line_spacing(p, WD_LINE_SPACING.MULTIPLE, 1.25)
            add_run(p, title_text, bold=True,
                    font_name_cn="黑体", font_name_en="Times New Roman",
                    font_size_pt=14)   # 四号
            i += 1
            continue

        if subsec_match:
            num, title_text = subsec_match.group(1), subsec_match.group(2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            set_line_spacing(p, WD_LINE_SPACING.MULTIPLE, 1.25)
            add_run(p, num + " " + title_text, bold=True,
                    font_name_cn="黑体", font_name_en="Times New Roman",
                    font_size_pt=12)   # 小四号黑体
            i += 1
            continue

        if section_match:
            num, title_text = section_match.group(1), section_match.group(2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            set_line_spacing(p, WD_LINE_SPACING.MULTIPLE, 1.25)
            add_run(p, num + " " + title_text, bold=True,
                    font_name_cn="黑体", font_name_en="Times New Roman",
                    font_size_pt=14)   # 四号黑体
            i += 1
            continue

        if chapter_match:
            num, title_text = chapter_match.group(1), chapter_match.group(2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            set_line_spacing(p, WD_LINE_SPACING.MULTIPLE, 1.25)
            add_run(p, num + ". " + title_text, bold=True,
                    font_name_cn="黑体", font_name_en="Times New Roman",
                    font_size_pt=15)   # 小三号黑体
            i += 1
            continue

        # ── 摘要/ABSTRACT 标题 ────────────────────────
        if re.match(r'^\*\*摘\s*要\*\*$', stripped) or stripped in ("摘　要", "摘要"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(12)
            add_run(p, "摘　要", bold=True,
                    font_name_cn="黑体", font_name_en="Times New Roman",
                    font_size_pt=18)   # 小二号
            i += 1
            continue

        if re.match(r'^\*\*ABSTRACT\*\*$', stripped) or stripped == "ABSTRACT":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(12)
            add_run(p, "ABSTRACT", bold=True,
                    font_name_cn="黑体", font_name_en="Times New Roman",
                    font_size_pt=18)
            i += 1
            continue

        # ── 目录标题 ──────────────────────────────────
        if re.match(r'^目\s*录$', stripped):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(12)
            add_run(p, "目　　录", bold=True,
                    font_name_cn="黑体", font_name_en="Times New Roman",
                    font_size_pt=18)
            i += 1
            continue

        # ── 目录条目（含页码，如 "绪论 1" "1.1 xxx 5"）────
        toc_match = re.match(r'^(.+?)\s+(\d+)$', stripped)
        if toc_match and not stripped.startswith('|'):
            entry_text = toc_match.group(1)
            page_num   = toc_match.group(2)
            # 判断是否是章级（不含小数点的数字开头，或特殊章名）
            is_chapter = bool(re.match(r'^(\d+\.|绪论|结论|参考文献|致谢|附录)', entry_text))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_line_spacing(p, WD_LINE_SPACING.MULTIPLE, 1.25)
            add_run(p, entry_text + "  " + page_num,
                    bold=is_chapter,
                    font_name_cn="黑体" if is_chapter else "宋体",
                    font_name_en="Times New Roman",
                    font_size_pt=12)
            i += 1
            continue

        # ── 关键词行 ──────────────────────────────────
        kw_cn = re.match(r'^关键词[：:](.+)', stripped)
        kw_en = re.match(r'^\*\*Key\s*Words?:\*\*(.+)', stripped, re.IGNORECASE) or \
                re.match(r'^Key\s*Words?:(.+)', stripped, re.IGNORECASE)
        if kw_cn:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_line_spacing(p, WD_LINE_SPACING.MULTIPLE, 1.25)
            add_run(p, "关键词：", bold=True,
                    font_name_cn="黑体", font_name_en="Times New Roman", font_size_pt=12)
            add_run(p, kw_cn.group(1).strip(),
                    font_name_cn="宋体", font_name_en="Times New Roman", font_size_pt=12)
            i += 1
            continue
        if kw_en:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_line_spacing(p, WD_LINE_SPACING.MULTIPLE, 1.25)
            add_run(p, "Key Words: ", bold=True,
                    font_name_cn="宋体", font_name_en="Times New Roman", font_size_pt=12)
            add_run(p, kw_en.group(1).strip(),
                    font_name_cn="宋体", font_name_en="Times New Roman", font_size_pt=12)
            i += 1
            continue

        # ── 参考文献条目 [数字] ────────────────────────
        ref_match = re.match(r'^\[(\d+)\]\s+(.+)', stripped)
        if ref_match:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent       = Cm(0.74)
            p.paragraph_format.first_line_indent = Cm(-0.74)
            set_line_spacing(p, WD_LINE_SPACING.MULTIPLE, 1.25)
            add_run(p, f"[{ref_match.group(1)}]  {ref_match.group(2)}",
                    font_name_cn="宋体", font_name_en="Times New Roman", font_size_pt=9)
            i += 1
            continue

        # ── "此处添加xxx图" 提示行 ─────────────────────
        if stripped.startswith("此处添加"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_line_spacing(p, WD_LINE_SPACING.MULTIPLE, 1.25)
            add_run(p, f"【{stripped}】",
                    font_name_cn="宋体", font_name_en="Times New Roman",
                    font_size_pt=12, color=RGBColor(0xFF, 0x00, 0x00))
            i += 1
            continue

        # ── 封面/声明页的粗体独立行 ───────────────────
        if re.match(r'^\*\*.+\*\*$', stripped):
            inner = stripped.strip('*')
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_line_spacing(p, WD_LINE_SPACING.MULTIPLE, 1.25)
            add_run(p, inner, bold=True,
                    font_name_cn="宋体", font_name_en="Times New Roman", font_size_pt=12)
            i += 1
            continue

        # ── 普通正文段落 ──────────────────────────────
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_line_spacing(p, WD_LINE_SPACING.MULTIPLE, 1.25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

        # 判断是否需要首行缩进（正文段落，非列表项、非图注、非封面行）
        is_list_item = re.match(r'^（\d+）|^\(\d+\)|^[①②③④⑤⑥⑦⑧⑨⑩]', stripped)
        is_cover_line = re.match(r'^(论文题目|姓名|学号|专业|指导教师|完成时间|签字日期|毕业设计|本人声明|学位论文)', stripped)
        if not is_list_item and not is_cover_line:
            set_first_indent(p, 24)   # 首行缩进2字符（小四号=12pt，2字符=24pt）

        add_inline_text(p, stripped,
                        font_name_cn="宋体", font_name_en="Times New Roman",
                        font_size_pt=12)
        i += 1

    doc.save(out_path)
    print(f"✅ 转换完成：{out_path}")


if __name__ == "__main__":
    convert(MD_FILE, OUT_FILE)
